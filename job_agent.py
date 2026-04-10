import streamlit as st
from openai import OpenAI
import io
import os
import sqlite3
import pandas as pd
import requests
from bs4 import BeautifulSoup
from datetime import datetime, timedelta
from docx import Document
from PyPDF2 import PdfReader
import json

# --- KONFIGURATION & DATABASE ---
st.set_page_config(page_title="Job Agent Pro - Master Suite", page_icon="🚀", layout="wide")
db_path = "job_agent_arkiv.db"

def get_danish_time():
    return (datetime.utcnow() + timedelta(hours=2)).strftime("%d. %m. %Y, %H:%M")

def init_db():
    conn = sqlite3.connect(db_path)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS archive
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  date TEXT, company TEXT, title TEXT, 
                  ansogning TEXT, opslag TEXT, tone TEXT)''')
    conn.commit()
    conn.close()

init_db()

# --- SESSION STATE ---
if 'step' not in st.session_state: st.session_state.step = 1
def next_step(): st.session_state.step += 1
def prev_step(): st.session_state.step -= 1
def reset(): 
    for key in list(st.session_state.keys()): del st.session_state[key]
    st.session_state.step = 1

# --- HJÆLPEFUNKTIONER ---
def get_text_from_url(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, 'html.parser')
        for script in soup(["script", "style"]): script.extract()
        return soup.get_text(separator=' ', strip=True)
    except: return ""

def extract_pdf(file):
    try:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    except: return ""

def fill_docx(template, replacements):
    try:
        template.seek(0)
        doc = Document(template)
        
        # Gennemgå alle afsnit
        for p in doc.paragraphs:
            for tag, content in replacements.items():
                if tag in p.text:
                    if "\n" in str(content):
                        # Håndtering af flere linjer (f.eks. erfaring eller ansøgning)
                        p.text = p.text.replace(tag, "")
                        cursor = p
                        for line in str(content).split('\n'):
                            if line.strip():
                                new_p = doc.add_paragraph(line.strip(), style=p.style)
                                cursor._element.addnext(new_p._element)
                                cursor = new_p
                    else:
                        # Simpel tekst erstatning
                        p.text = p.text.replace(tag, str(content))

        # Gennemgå alle tabeller (vigtigt for CV-layouts)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for tag, content in replacements.items():
                            if tag in p.text:
                                p.text = p.text.replace(tag, str(content))
        
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    except Exception as e:
        st.error(f"Fejl i Word-fletning: {e}")
        return None

# --- APP FLOW ---
st.title("💼 Job Agent Pro")
st.progress(st.session_state.step / 4)

if st.session_state.step == 1:
    st.header("1. Upload dine dokumenter")
    cv_file = st.file_uploader("Upload dit Master-CV (PDF)", type="pdf")
    col_t1, col_t2 = st.columns(2)
    temp_ans = col_t1.file_uploader("Din Word-skabelon: Ansøgning", type="docx")
    temp_cv = col_t2.file_uploader("Din Word-skabelon: CV", type="docx")
    
    c1, c2, c3 = st.columns(3)
    comp = c1.text_input("Virksomhedens navn:")
    titl = c2.text_input("Jobtitel:")
    name = c3.text_input("Dit fulde navn:")
    contact_person = st.text_input("Kontaktperson hos virksomheden:")
    
    if st.button("Næste →", disabled=not (cv_file and comp and titl and name)):
        st.session_state.cv_text = extract_pdf(cv_file)
        st.session_state.temp_ans = temp_ans
        st.session_state.temp_cv = temp_cv
        st.session_state.comp = comp
        st.session_state.titl = titl
        st.session_state.name = name
        st.session_state.contact = contact_person
        next_step()
        st.rerun()

elif st.session_state.step == 2:
    st.header("2. Jobopslaget")
    url = st.text_input("Indsæt link til jobopslag:")
    if st.button("Hent tekst fra link") and url:
        txt = get_text_from_url(url)
        if txt: st.session_state.fetched_txt = txt
    
    opslag = st.text_area("Jobopslagets tekst:", value=st.session_state.get('fetched_txt', ""), height=300)
    noter = st.text_area("Eventuelle noter (hvad vil du fremhæve?):")
    
    col1, col2 = st.columns(2)
    if col1.button("← Tilbage"): prev_step(); st.rerun()
    if col2.button("Generér Analyse & Strategi →", disabled=not opslag):
        st.session_state.opslag = opslag
        st.session_state.noter = noter
        next_step()
        st.rerun()

elif st.session_state.step == 3:
    st.header("3. Tilpas strategi")
    c1, c2 = st.columns(2)
    tone = c1.selectbox("Tone og sprogstil:", ["Professionel", "Balanceret", "Personlig", "Kreativ", "Formel"])
    h_type = c2.selectbox("Overskriftstype:", ["Værdiskabende", "Formel", "Catchy", "Spørgende"])
    
    f1, f2 = st.columns(2)
    fokus = f1.radio("Fokus i ansøgning:", ["Faglige resultater", "Personlige kompetencer", "Balanceret"], horizontal=True)
    mot_pos = f2.radio("Motivationens placering:", ["I starten (krogen)", "I bunden (opsamlingen)"])
    
    if st.button("Generér Alt (Ansøgning + CV) ✨"):
        st.session_state.p = {"tone": tone, "h_type": h_type, "fokus": fokus, "mot_pos": mot_pos}
        next_step()
        st.rerun()

elif st.session_state.step == 4:
    st.header("4. Resultat & Download")
    if "final_res" not in st.session_state:
        with st.spinner("AI optimerer nu din profil og skriver dokumenterne..."):
            try:
                client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
                p = st.session_state.p
                
                # ATS Keyword Analyse
                ats_resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": f"Find de 10 vigtigste keywords i dette opslag:\n{st.session_state.opslag}"}]
                )
                st.session_state.ats_result = ats_resp.choices[0].message.content

                # Hovedgenerering
                main_prompt = f"""
                Lav en JSON pakke på dansk. Alt indhold skal matche tonen '{p['tone']}'.
                Brug direkte vendinger og keywords fra jobopslaget for at sikre ATS-match.
                
                1. 'ansogning': En komplet, overbevisende ansøgning (min. 5-6 afsnit). Motivation skal placeres {p['mot_pos']}.
                2. 'overskrift': En stærk overskrift (type: {p['h_type']}).
                3. 'cv_profil': En målrettet profiltekst til CV'et.
                4. 'cv_erfaring': Optimeret beskrivelse af tidligere erfaringer. Brug keywords fra jobopslaget til at beskrive ansvarsområder og resultater.
                5. 'cv_kompetencer': Liste over relevante kompetencer (Hard & Soft skills).
                6. 'cv_uddannelse': Uddannelseshistorik fra CV data.
                7. 'cv_kurser': Kurser og certificeringer.
                8. 'cv_sprog': Sprogkundskaber.
                9. 'pitch': 3-4 sætninger til LinkedIn.
                10. 'interview': 3 kritiske spørgsmål og svar i Markdown.

                DATA: CV: {st.session_state.cv_text}, JOB: {st.session_state.opslag}, NOTER: {st.session_state.noter}
                """
                
                resp = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "system", "content": "Du er elite-rekrutteringskonsulent. Svar KUN i JSON."}, 
                             {"role": "user", "content": main_prompt}],
                    response_format={"type": "json_object"}
                )
                st.session_state.final_res = json.loads(resp.choices[0].message.content)

                # Arkiv gemning
                conn = sqlite3.connect(db_path); c = conn.cursor()
                c.execute("INSERT INTO archive (date, company, title, ansogning, opslag, tone) VALUES (?,?,?,?,?,?)",
                          (get_danish_time(), st.session_state.comp, st.session_state.titl, st.session_state.final_res['ansogning'], st.session_state.opslag, p['tone']))
                conn.commit(); conn.close()
            except Exception as e:
                st.error(f"Fejl: {e}")

    if "final_res" in st.session_state:
        res = st.session_state.final_res
        st.info(f"**ATS Keyword Fokus:**\n{st.session_state.ats_result}")
        
        col_left, col_right = st.columns(2)
        
        with col_left:
            st.subheader("📄 Ansøgning")
            st.markdown(f"**{res.get('overskrift')}**")
            st.write(res.get('ansogning'))
            
            if st.session_state.temp_ans:
                ans_data = {
                    "{{ANSOGNING}}": res.get('ansogning'),
                    "{{OVERSKRIFT}}": res.get('overskrift'),
                    "{{VIRKSOMHED}}": st.session_state.comp,
                    "{{JOBTITEL}}": st.session_state.titl,
                    "{{KONTAKTPERSON}}": st.session_state.contact,
                    "{{DATO}}": datetime.now().strftime("%d. %m. %Y")
                }
                ans_doc = fill_docx(st.session_state.temp_ans, ans_data)
                st.download_button("Download Ansøgning (.docx)", ans_doc, f"Ansøgning_{st.session_state.comp}.docx")

        with col_right:
            st.subheader("👤 Optimeret CV")
            st.write(res.get('cv_profil'))
            st.divider()
            st.write(res.get('cv_erfaring'))
            
            if st.session_state.temp_cv:
                cv_data = {
                    "{{NAVN}}": st.session_state.name,
                    "{{CV_PROFIL}}": res.get('cv_profil'),
                    "{{CV_ERFARING}}": res.get('cv_erfaring'),
                    "{{CV_KOMPETENCER}}": res.get('cv_kompetencer'),
                    "{{CV_UDDANNELSE}}": res.get('cv_uddannelse'),
                    "{{CV_KURSER}}": res.get('cv_kurser'),
                    "{{CV_SPROG}}": res.get('cv_sprog'),
                    "{{JOBTITEL}}": st.session_state.titl,
                    "{{VIRKSOMHED}}": st.session_state.comp,
                    "{{DATO}}": datetime.now().strftime("%d. %m. %Y")
                }
                cv_doc = fill_docx(st.session_state.temp_cv, cv_data)
                st.download_button("Download Optimeret CV (.docx)", cv_doc, f"CV_{st.session_state.comp}.docx")

        st.divider()
        c1, c2 = st.columns(2)
        with c1:
            st.subheader("✉️ LinkedIn Pitch")
            st.success(res.get('pitch'))
        with c2:
            st.subheader("🎤 Interviewforberedelse")
            st.markdown(res.get('interview'))
        
        if st.button("Start forfra 🔄"): reset(); st.rerun()

# --- ARKIV ---
st.divider()
st.subheader("📂 Tidligere ansøgninger")
if os.path.exists(db_path):
    conn = sqlite3.connect(db_path); df = pd.read_sql_query("SELECT * FROM archive ORDER BY id DESC LIMIT 10", conn); conn.close()
    for i, row in df.iterrows():
        with st.expander(f"📌 {row['company']} - {row['title']} ({row['date']})"):
            st.write(row['ansogning'])
